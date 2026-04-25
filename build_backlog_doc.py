from __future__ import annotations

import csv
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\dev\workspace\evolucaoFisica")
PRODUCT_DIR = ROOT / "arquivos" / "produto"
CSV_PATH = PRODUCT_DIR / "backlog_priorizado.csv"
OUTPUT_PATH = PRODUCT_DIR / "backlog_priorizado_exportavel.docx"

def style_text(paragraph, *, size=10.5, bold=False, color=None, align=None):
    if align is not None:
        paragraph.alignment = align
    for run in paragraph.runs:
        run.font.name = "Aptos"
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor.from_string(color)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.add_run(text)
    style_text(paragraph, size=10.5)


def add_body(document: Document, text: str, *, size=10.5) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run(text)
    style_text(paragraph, size=size)


def normalize(text: str | None) -> str:
    return text.strip() if text and text.strip() else "-"


rows = []
with CSV_PATH.open("r", encoding="utf-8") as handle:
    reader = csv.DictReader(handle)
    rows = list(reader)

doc = Document()

section = doc.sections[0]
section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run("Backlog Priorizado e Mapa Funcional")
title_run.bold = True
title_run.font.name = "Aptos Display"
title_run.font.size = Pt(22)
title_run.font.color.rgb = RGBColor.from_string("123B5D")

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run("Projeto Evolucao Fisica | Visao executiva para backend, mobile e produto")
subtitle_run.font.name = "Aptos"
subtitle_run.font.size = Pt(11)
subtitle_run.font.color.rgb = RGBColor.from_string("5C6B73")

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_run = meta.add_run(f"Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')}")
meta_run.font.name = "Aptos"
meta_run.font.size = Pt(9)
meta_run.font.color.rgb = RGBColor.from_string("6B7280")

doc.add_paragraph("")

summary_heading = doc.add_paragraph()
summary_heading.add_run("Resumo Executivo").bold = True
style_text(summary_heading, size=14, color="123B5D")

for bullet in [
    "Separar plano/modelo de execucao real e obrigatorio para historico, gamificacao e analytics confiaveis.",
    "A prioridade do produto deve ser fundacao tecnica, treino e alimentacao reais, motor de gamificacao e so depois social expandido.",
    "As regras de recompensa precisam continuar baseadas apenas em fatos persistidos e auditaveis.",
]:
    add_bullet(doc, bullet)

doc.add_paragraph("")

phase_heading = doc.add_paragraph()
phase_heading.add_run("Ordem Recomendada de Entrega").bold = True
style_text(phase_heading, size=14, color="123B5D")

for item in [
    "Fase 1: identidade, onboarding, metas do atleta e base tecnica de producao.",
    "Fase 2: treino planejado + treino real como fato gerador central.",
    "Fase 3: alimentacao com base nutricional, plano semanal e refeicao real.",
    "Fase 4: gamificacao auditavel sobre eventos reais.",
    "Fase 5: social com privacidade, grupos e ranking contextual.",
    "Fase 6: notificacoes, analytics, integracoes e experiencia offline.",
]:
    add_bullet(doc, item)

doc.add_paragraph("")

table_heading = doc.add_paragraph()
table_heading.add_run("Backlog Priorizado por Faixa").bold = True
style_text(table_heading, size=14, color="123B5D")

priority_descriptions = {
    "P0": "Fundacao obrigatoria para o produto operar com seguranca e consistencia.",
    "P1": "Escopo principal do produto para treino, alimentacao e gamificacao funcionarem de ponta a ponta.",
    "P2": "Capacidades de retencao, social e analytics que elevam a experiencia.",
    "P3": "Expansoes de experiencia, integracoes e recursos avancados.",
}

for priority in ["P0", "P1", "P2", "P3"]:
    section_rows = [row for row in rows if row["prioridade"] == priority]
    heading = doc.add_paragraph()
    heading.add_run(f"{priority}").bold = True
    style_text(heading, size=12.5, color="123B5D")

    add_body(doc, priority_descriptions[priority], size=10)

    for row in section_rows:
        item_heading = doc.add_paragraph()
        item_heading.add_run(f"{row['id']} | {row['modulo']} | {row['feature']}").bold = True
        style_text(item_heading, size=10.5, color="1F2937")

        add_bullet(doc, f"Epic: {normalize(row['epic'])}")
        add_bullet(doc, f"Tipo: {normalize(row['tipo'])}")
        add_bullet(doc, f"Status sugerido: {normalize(row['status_sugerido'])}")
        add_bullet(doc, f"Dependencias: {normalize(row['dependencias'])}")
        doc.add_paragraph("")

doc.add_page_break()

risks_heading = doc.add_paragraph()
risks_heading.add_run("Riscos Tecnicos Principais").bold = True
style_text(risks_heading, size=14, color="123B5D")

for risk in [
    "Misturar plano com execucao real de treino e dieta, contaminando historico, analytics e recompensa.",
    "Manter gamificacao sem mecanismos robustos de reprocessamento e idempotencia.",
    "Escalar feed, ranking e agregacoes sem paginacao e indices coerentes.",
    "Iniciar Flutter sem contratos consolidados para estados de execucao e agregados.",
    "Persistir sem estrategia formal de migracao de banco e seguranca de identidade.",
]:
    add_bullet(doc, risk)

doc.save(OUTPUT_PATH)
print(OUTPUT_PATH)
